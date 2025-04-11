import streamlit as st
from typing import List
import os
import docx
import io

from PyPDF2 import PdfReader
from langchain_core.prompts import ChatPromptTemplate
from langchain_community.llms import Ollama
from langchain_ollama import OllamaEmbeddings
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_core.output_parsers import StrOutputParser
from langchain_community.vectorstores import Chroma
from langchain_core.runnables import RunnablePassthrough
from langchain_core.documents import Document
from langchain_core.messages import AIMessage, HumanMessage


# Initialize session state
def init_session_state():
    if "chat_history" not in st.session_state:
        st.session_state.chat_history = []
    if "retriever" not in st.session_state:
        st.session_state.retriever = None
    if "conversation" not in st.session_state:
        st.session_state.conversation = None

# Define LLM model
@st.cache_resource
def get_llm():
    return Ollama(model="deepseek-r1:14b")

# Define embeddings
@st.cache_resource
def get_embeddings():
    return OllamaEmbeddings(model="nomic-embed-text:latest")

# Define prompt template
template = """
You are an AI assistant for question-answering tasks, designed to provide responses based on the retrieved context. Your goal is to generate an answer that matches the level of detail required by the question.

- If the question is broad, complex, or requires explanation, provide a detailed and well-structured response with examples where applicable.
- If the question is straightforward or factual, provide a concise response.
- If the context does not contain the answer, say that you don't know instead of making up information.

Context: {context}  
Question: {question}  

Answer: 
"""

# Function to extract text from different file types
def extract_text(uploaded_file) -> str:
    if uploaded_file is None:
        return ""
    
    file_extension = uploaded_file.name.split('.')[-1].lower()
    
    try:
        if file_extension == 'pdf':
            pdf_reader = PdfReader(uploaded_file)
            text = ""
            for page in pdf_reader.pages:
                text += page.extract_text() or ""
            return text.strip()
            
        elif file_extension == 'txt':
            return uploaded_file.getvalue().decode('utf-8').strip()
            
        elif file_extension in ['doc', 'docx']:
            doc = docx.Document(io.BytesIO(uploaded_file.getvalue()))
            text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
            return text.strip()
            
        else:
            raise ValueError(f"Unsupported file type: {file_extension}")
            
    except Exception as e:
        raise Exception(f"Error processing {file_extension.upper()} file: {str(e)}")

# Function to split text into chunks
def split_text(text: str) -> List[Document]:
    splitter = RecursiveCharacterTextSplitter(
        chunk_size=1000,
        chunk_overlap=200,
        separators=["\n\n", "\n", ".", "!", "?", ",", " ", ""],
    )
    chunks = splitter.split_text(text)
    return [Document(page_content=chunk) for chunk in chunks]

# Function to create a retriever from the vector database
def create_retriever(documents: List[Document]):
    embeddings = get_embeddings()
    persistent_directory = "chroma_db"
    vector_store = Chroma.from_documents(
        documents=documents,
        embedding=embeddings,
        persist_directory=persistent_directory
    )
    
    return vector_store.as_retriever(search_kwargs={"k": 3})

# Function to create conversation chain
def create_conversation_chain(retriever):
    prompt = ChatPromptTemplate.from_template(template)
    llm = get_llm()

    def format_context(inputs):
        docs = retriever.invoke(inputs["question"])
        return {
            **inputs,
            "context": "\n\n".join([doc.page_content for doc in docs])
        }

    chain = (
        RunnablePassthrough() 
        | format_context
        | prompt
        | llm
        | StrOutputParser()
    )
    
    return chain
def main():
    st.title("Local AI Chat")
    init_session_state()

    # Sidebar for file upload
    with st.sidebar:
        st.header("Document Upload")
        uploaded_files = st.file_uploader(
            "Upload documents", 
            type=["pdf", "txt", "doc", "docx"],
            accept_multiple_files=True
        )
        
        if uploaded_files and st.button("Process Documents"):
            with st.spinner("Processing documents..."):
                try:
                    # Extract and process text from all documents
                    combined_text = ""
                    for file in uploaded_files:
                        text = extract_text(file)
                        if text:
                            combined_text += f"\n\n{text}"
                    
                    if not combined_text:
                        st.error("No text could be extracted from the documents. Please try other files.")
                        return
                    
                    # Create documents and retriever
                    documents = split_text(combined_text)
                    st.session_state.retriever = create_retriever(documents)
                    st.session_state.conversation = create_conversation_chain(st.session_state.retriever)
                    st.success(f"Documents processed successfully! You can now ask questions about the content.")
                except Exception as e:
                    st.error(f"An error occurred while processing the documents: {str(e)}")
                    return

    # Main chat interface
    if st.session_state.conversation is None:
        st.info("Please upload and process documents to start asking questions.")
        return

    # Display chat history
    for message in st.session_state.chat_history:
        with st.chat_message("Human" if isinstance(message, HumanMessage) else "AI"):
            st.markdown(message.content)

    # Chat input
    if query := st.chat_input("Ask a question about your documents"):
        # Add user message to chat history
        st.session_state.chat_history.append(HumanMessage(content=query))
        
        # Display user message
        with st.chat_message("Human"):
            st.markdown(query)

        # Generate and display AI response
        with st.chat_message("AI"):
            response = ""
            response_container = st.empty()
            
            for chunk in st.session_state.conversation.stream({"question": query}):
                response += chunk
                response_container.markdown(response)
            
            # Add AI response to chat history
            st.session_state.chat_history.append(AIMessage(content=response))

if name == "__main__":
    main()
